from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.http import HttpResponse
from django.template import loader
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.shortcuts import render
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import pytesseract
from PIL import Image
import io
from django.core.exceptions import ValidationError
import fitz  # PyMuPDF
import docx
from docx import Document  # Ensure you have `python-docx` installed

# Specify the path to the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def validate_file(file):
    valid_mime_types = [
        'image/jpeg', 'image/png',  # For images
        'application/pdf',  # For PDFs
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # For DOCX
        'application/msword'  # For DOC
    ]
    if file.content_type not in valid_mime_types:
        raise ValidationError('Unsupported file type.')

def extract_text_from_pdf(file_stream):
    pdf_document = fitz.open(stream=file_stream, filetype='pdf')
    text = ""
    for page in pdf_document:
        text += page.get_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file_stream):
    extracted_text = ''
    try:
        document = Document(file_stream)
        for paragraph in document.paragraphs:
            extracted_text += paragraph.text + '\n'
    except Exception as e:
        extracted_text = f"Error reading DOCX: {e}"
    return extracted_text

def index(request):
    return render(request, 'index.html')

import mimetypes

#image
@csrf_exempt
def upload_image(request):
    extracted_text = ''
    uploaded_file_url = ''
    file_content = b''

    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        
        try:
            # Read the file content
            file_content = uploaded_file.read()
            file_stream = io.BytesIO(file_content)
            
            # Attempt to get MIME type from the uploaded file
            mime_type = uploaded_file.content_type
            print(f"Received MIME type: {mime_type}")

            # Fallback to mimetypes library if content_type is missing
            if not mime_type:
                mime_type, _ = mimetypes.guess_type(uploaded_file.name)
                print(f"Guessed MIME type using mimetypes: {mime_type}")
            
            # Process the file based on its MIME type
            if mime_type in ['image/jpeg', 'image/png']:
                image = Image.open(file_stream)
                image.verify()
                file_stream.seek(0)
                extracted_text = pytesseract.image_to_string(Image.open(file_stream))
            else:
                extracted_text = "Unsupported file type."
            
        except Exception as e:
            extracted_text = f"Error processing file: {e}"

        if file_content:
            file_name = default_storage.save(uploaded_file.name, ContentFile(file_content))
            uploaded_file_url = default_storage.url(file_name)

        if request.headers.get('Accept') == 'application/json':
            return JsonResponse({'extracted_text': extracted_text})
    
    return render(request, 'upload_image.html', {
        'uploaded_file_url': uploaded_file_url,
        'extracted_text': extracted_text
    })

from PyPDF2 import PdfReader  # Make sure PyPDF2 is installed
#pdf
@csrf_exempt
def upload_pdf_file(request):
    extracted_text = ''
    uploaded_file_url = ''
    file_content = b''

    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        
        try:
            # Read the file content
            file_content = uploaded_file.read()
            file_stream = io.BytesIO(file_content)
            
            # Attempt to get MIME type from the uploaded file
            mime_type = uploaded_file.content_type
            print(f"Received MIME type: {mime_type}")

            # Fallback to mimetypes library if content_type is missing
            if not mime_type:
                mime_type, _ = mimetypes.guess_type(uploaded_file.name)
                print(f"Guessed MIME type using mimetypes: {mime_type}")
            
            # Process the file based on its MIME type
            if mime_type == 'application/pdf':
                # Extract text from PDF
                extracted_text = extract_text_from_pdf(file_stream)
            else:
                extracted_text = "Unsupported file type."
            
        except Exception as e:
            extracted_text = f"Error processing file: {e}"

        if file_content:
            file_name = default_storage.save(uploaded_file.name, ContentFile(file_content))
            uploaded_file_url = default_storage.url(file_name)

        if request.headers.get('Accept') == 'application/json':
            return JsonResponse({'uploaded_file_url': uploaded_file_url, 'extracted_text': extracted_text})
    
    return render(request, 'upload_pdf.html', {
        'uploaded_file_url': uploaded_file_url,
        'extracted_text': extracted_text
    })

def extract_text_from_pdf(file_stream):
    extracted_text = ''
    try:
        pdf_reader = PdfReader(file_stream)
        for page in pdf_reader.pages:
            extracted_text += page.extract_text()
    except Exception as e:
        extracted_text = f"Error reading PDF: {e}"
    return extracted_text

#docx.
@csrf_exempt
def upload_docx_file(request):
    extracted_text = ''
    uploaded_file_url = ''
    file_content = b''

    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        
        try:
            # Read the file content
            file_content = uploaded_file.read()
            file_stream = io.BytesIO(file_content)
            
            # Attempt to get MIME type from the uploaded file
            mime_type = uploaded_file.content_type
            print(f"Received MIME type: {mime_type}")

            # Fallback to mimetypes library if content_type is missing
            if not mime_type:
                mime_type, _ = mimetypes.guess_type(uploaded_file.name)
                print(f"Guessed MIME type using mimetypes: {mime_type}")
            
            # Process the file based on its MIME type
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                extracted_text = extract_text_from_docx(file_stream)
            else:
                extracted_text = "Unsupported file type."
            
        except Exception as e:
            extracted_text = f"Error processing file: {e}"

        if file_content:
            file_name = default_storage.save(uploaded_file.name, ContentFile(file_content))
            uploaded_file_url = default_storage.url(file_name)

        if request.headers.get('Accept') == 'application/json':
            return JsonResponse({'uploaded_file_url': uploaded_file_url, 'extracted_text': extracted_text})
    
    return render(request, 'upload_docx.html', {
        'uploaded_file_url': uploaded_file_url,
        'extracted_text': extracted_text
    })

@csrf_exempt
def download_text(request):
    format = request.GET.get('format', 'text')
    text = request.GET.get('text', '')

    # Debug: Print the text content to the console
    print(f"Text content to download: {text}")

    if format == 'text':
        response = HttpResponse(text, content_type='text/plain')
        response['Content-Disposition'] = 'attachment; filename="extracted_text.txt"'
        return response

    elif format == 'pdf':
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        text_object = p.beginText(40, 750)
        text_object.setFont("Helvetica", 12)
        for line in text.split('\n'):
            text_object.textLine(line)
        p.drawText(text_object)
        p.showPage()
        p.save()
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="extracted_text.pdf"'
        return response

    elif format == 'docx':
        buffer = BytesIO()
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="extracted_text.docx"'
        return response

    return HttpResponse("Invalid format", status=400)
